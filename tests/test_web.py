"""Prueba el flujo HTTP completo: login, subida, vista previa y descarga."""
from __future__ import annotations

import importlib
import io
import os

import pytest
from docx import Document

from tests import fixtures

PASSWORD = "clave-de-prueba"


@pytest.fixture(scope="module")
def client():
    os.environ["APP_PASSWORD"] = PASSWORD
    os.environ["SECRET_KEY"] = "secreto-para-tests"
    os.environ["COOKIE_SECURE"] = "0"
    from fastapi.testclient import TestClient

    from app import auth, config, main, store

    importlib.reload(config)
    importlib.reload(auth)
    importlib.reload(store)
    importlib.reload(main)
    return TestClient(main.app)


@pytest.fixture(scope="module")
def logged_in(client):
    response = client.post("/login", data={"password": PASSWORD}, follow_redirects=False)
    assert response.status_code == 303
    return client


def test_health_reports_password_configured(client):
    assert client.get("/health").json() == {"status": "ok", "password_configured": True}


def test_home_requires_the_password(client):
    fresh = client.__class__(client.app)
    body = fresh.get("/").text
    assert "Clave de acceso" in body and "Cargar reportes" not in body


def test_wrong_password_is_rejected(client):
    fresh = client.__class__(client.app)
    response = fresh.post("/login", data={"password": "incorrecta"})
    assert response.status_code == 401
    assert "Clave incorrecta" in response.text


def test_api_rejects_anonymous_requests(client):
    fresh = client.__class__(client.app)
    assert fresh.post("/api/preview", json={"job_id": "x"}).status_code == 401


@pytest.fixture(scope="module")
def job(logged_in):
    files = {
        "sag": ("sag.xlsx", fixtures.sag_xlsx(), "application/vnd.ms-excel"),
        "cable": ("cable.xlsx", fixtures.cable_xlsx(), "application/vnd.ms-excel"),
        "structures": ("estructuras.xlsx", fixtures.structures_xlsx(), "application/vnd.ms-excel"),
    }
    response = logged_in.post("/api/upload", files=files)
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["ready"] is True, payload["error"]
    return payload


def test_upload_detects_every_required_column(job):
    for report in job["reports"].values():
        assert report["missing"] == []


def test_upload_lists_cables_and_structures(job):
    values = sorted(round(c["value"], 4) for c in job["options"]["cables"])
    assert values == [fixtures.CABLE_A, fixtures.CABLE_B]
    kinds = {s["key"]: s["kind"] for s in job["options"]["structures"]}
    assert kinds == {"5": "anclaje", "6": "anclaje", "7": "suspension", "8": "anclaje"}


def test_upload_reports_weather_cases_and_condition(job):
    assert job["options"]["condition_text"] == "Initial RS"
    assert "Hielo" in job["options"]["weather_cases"]
    assert job["options"]["has_sections"] is True


def _config(job):
    return {
        "job_id": job["job_id"],
        "cable": fixtures.CABLE_A,
        "temperatures": [float(t) for t in fixtures.TEMPS],
        "kinds": {s["key"]: s["kind"] for s in job["options"]["structures"]},
        "prefix": "E",
    }


def test_preview_lists_sections_with_kinds_vanos_and_cable(logged_in, job):
    data = logged_in.post("/api/preview", json=_config(job)).json()
    assert [s["tramo"] for s in data["sections"]] == ["E5-E6", "E6-E8"]
    simple, with_suspension = data["sections"]
    assert simple["from_kind"] == "anclaje" and simple["to_kind"] == "anclaje"
    assert round(simple["subspans"][0]["vano"], 1) == 88.1
    assert [i["label"] for i in with_suspension["intermediate"]] == ["E7"]
    assert [round(s["vano"], 1) for s in with_suspension["subspans"]] == [86.3, 80.9]
    assert simple["cable"] == pytest.approx(fixtures.CABLE_A)


def test_reclassifying_a_structure_changes_the_sections(logged_in, job):
    payload = _config(job)
    payload["kinds"]["7"] = "anclaje"  # E7 pasa a ser anclaje
    data = logged_in.post("/api/preview", json=payload).json()
    assert [s["tramo"] for s in data["sections"]] == ["E5-E6", "E6-E7", "E7-E8"]


def test_options_endpoint_switches_cables_with_the_weather_case(logged_in, job):
    body = {"job_id": job["job_id"], "cable": None, "weather_case": "Hielo"}
    data = logged_in.post("/api/options", json=body).json()
    values = sorted(round(c["value"], 4) for c in data["cables"])
    assert values == [round(fixtures.CABLE_A * fixtures.ICE_FACTOR, 4),
                      round(fixtures.CABLE_B * fixtures.ICE_FACTOR, 4)]
    assert data["temperatures"] == [float(t) for t in fixtures.TEMPS]


def test_preview_returns_structures_for_inline_editing(logged_in, job):
    data = logged_in.post("/api/preview", json=_config(job)).json()
    keys = {s["key"] for s in data["structures"]}
    assert keys == {"5", "6", "7", "8"}
    assert data["sections"][1]["intermediate"][0]["key"] == "7"


def test_generate_returns_a_word_document_with_both_tables(logged_in, job):
    payload = {**_config(job), "start_number": 4, "condicion_texto": "Initial RS"}
    response = logged_in.post("/api/generate", json=payload)
    assert response.status_code == 200
    assert response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "attachment; filename=" in response.headers["content-disposition"]
    document = Document(io.BytesIO(response.content))
    titles = [p.text for p in document.paragraphs if p.text.startswith("Tabla ")]
    assert titles == [
        "Tabla 4: Tramo entre las estructuras N°5 y N°6 en condición Initial RS",
        "Tabla 5: Tramo entre las estructuras N°6 y N°8 en condición Initial RS",
    ]
    assert len(document.tables) == 2


def test_selecting_fewer_temperatures_narrows_the_table(logged_in, job):
    payload = {**_config(job), "temperatures": [0.0, 20.0, 40.0]}
    document = Document(io.BytesIO(logged_in.post("/api/generate", json=payload).content))
    header = [c.text for c in document.tables[0].rows[1].cells]
    assert header[6:] == ["0°C", "20°C", "40°C"]


def test_expired_job_returns_a_clear_message(logged_in, job):
    response = logged_in.post("/api/preview", json={**_config(job), "job_id": "inexistente"})
    assert response.status_code == 404
    assert "vencio" in response.json()["detail"]


# --------------------------------------------------------------------------
# Seleccion de tramos
# --------------------------------------------------------------------------
def test_preview_gives_every_section_a_stable_key(logged_in, job):
    data = logged_in.post("/api/preview", json=_config(job)).json()
    claves = [s["key"] for s in data["sections"]]
    assert claves == ["5>6", "6>8"]
    # La clave no cambia al repetir la consulta.
    otra = logged_in.post("/api/preview", json=_config(job)).json()
    assert [s["key"] for s in otra["sections"]] == claves


def test_generating_only_the_checked_sections(logged_in, job):
    payload = {**_config(job), "sections": ["6>8"], "start_number": 1, "condicion_texto": "Initial RS"}
    document = Document(io.BytesIO(logged_in.post("/api/generate", json=payload).content))
    assert len(document.tables) == 1
    titulos = [p.text for p in document.paragraphs if p.style.name == "Caption"]
    assert titulos == ["Tabla 1: Tramo entre las estructuras N°6 y N°8 en condición Initial RS"]


def test_an_empty_selection_means_every_section(logged_in, job):
    payload = {**_config(job), "sections": []}
    document = Document(io.BytesIO(logged_in.post("/api/generate", json=payload).content))
    assert len(document.tables) == 2


def test_keys_that_no_longer_exist_are_rejected(logged_in, job):
    payload = {**_config(job), "sections": ["99>100"]}
    response = logged_in.post("/api/generate", json=payload)
    assert response.status_code == 400
    assert "marcados" in response.json()["detail"]


def test_section_keys_survive_reclassifying_a_structure(logged_in, job):
    """Marcar E5-E6 debe seguir valiendo aunque cambie el tipo de otra estructura."""
    payload = _config(job)
    payload["kinds"]["7"] = "anclaje"
    data = logged_in.post("/api/preview", json=payload).json()
    claves = [s["key"] for s in data["sections"]]
    assert claves == ["5>6", "6>7", "7>8"]

    payload["sections"] = ["5>6"]
    payload["start_number"] = 1
    payload["condicion_texto"] = "Initial RS"
    document = Document(io.BytesIO(logged_in.post("/api/generate", json=payload).content))
    assert len(document.tables) == 1
    titulos = [p.text for p in document.paragraphs if p.style.name == "Caption"]
    assert "N°5 y N°6" in titulos[0]
