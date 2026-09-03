"""Verifica que el flujo completo reproduzca las tablas 10-4 y 10-5 de referencia."""
from __future__ import annotations

import io

import pytest
from docx import Document
from docx.oxml.ns import qn

from app import analysis, docx_writer, parsing
from tests import fixtures


@pytest.fixture(scope="module")
def dataset() -> analysis.Dataset:
    sag_sheet = parsing.read_report(fixtures.sag_xlsx(), "sag.xlsx", "sag")
    cable_sheet = parsing.read_report(fixtures.cable_xlsx(), "cable.xlsx", "cable")
    struct_sheet = parsing.read_report(fixtures.structures_xlsx(), "str.xlsx", "structures")
    return analysis.build_dataset(
        sag_sheet, parsing.auto_map(sag_sheet, "sag"),
        cable_sheet, parsing.auto_map(cable_sheet, "cable"),
        struct_sheet, parsing.auto_map(struct_sheet, "structures"),
    )


def test_header_detection_finds_columns_despite_title_rows():
    sheet = parsing.read_report(fixtures.sag_xlsx(), "sag.xlsx", "sag")
    mapping = parsing.auto_map(sheet, "sag")
    assert mapping.missing_required() == []
    # Los espacios repetidos del export se colapsan al normalizar el encabezado.
    assert mapping.mapping["horz_tension"] == "Horz. Tension (daN)"
    assert mapping.mapping["temp"] == "Temp. (deg C)"
    assert mapping.mapping["span_from_str"] == "Span From Str."


def test_structures_are_classified_from_the_stk_filename():
    # La marca se busca en el nombre del archivo, no en la ruta: una carpeta
    # llamada "a" no debe leerse como _A_.
    assert analysis.classify_from_name(fixtures.NAMES["6"]) == analysis.ANCLAJE
    assert analysis.classify_from_name(fixtures.NAMES["7"]) == analysis.SUSPENSION
    assert analysis.classify_from_name(r"C:\Users\a\x\Portal_15.stk") is None


def test_stk_filename_wins_over_a_misleading_description(dataset):
    # Todas las estructuras traen "Portal anclaje 15m" como descripcion, pero
    # la 7 es _S_ en el nombre del archivo.
    assert dataset.structures["7"].auto_kind == analysis.SUSPENSION
    assert dataset.structures["6"].auto_kind == analysis.ANCLAJE


def test_condition_is_read_from_the_report(dataset):
    assert dataset.condition_text == "Initial RS"


def test_sections_come_from_the_pls_cadd_section_number(dataset):
    assert dataset.has_sections
    assert dataset.span_sections[("6", "7")] == dataset.span_sections[("7", "8")]
    assert dataset.span_sections[("5", "6")] != dataset.span_sections[("6", "7")]


def test_cable_options_ignore_the_ice_case_by_default(dataset):
    # Con hielo la carga vertical sube; el peso propio es el que identifica al cable.
    values = sorted(round(float(o["value"]), 4) for o in analysis.cable_options(dataset))
    assert values == [fixtures.CABLE_A, fixtures.CABLE_B]


def test_cable_options_follow_the_chosen_weather_case(dataset):
    values = sorted(round(float(o["value"]), 4) for o in analysis.cable_options(dataset, "Hielo"))
    assert values == [round(fixtures.CABLE_A * fixtures.ICE_FACTOR, 4),
                      round(fixtures.CABLE_B * fixtures.ICE_FACTOR, 4)]


def test_temperatures_match_the_report(dataset):
    assert analysis.temperature_options(dataset, fixtures.CABLE_A) == [float(t) for t in fixtures.TEMPS]


@pytest.fixture(scope="module")
def sections(dataset) -> list[analysis.Section]:
    kinds = {k: s.auto_kind for k, s in dataset.structures.items()}
    built, _ = analysis.build_sections(
        dataset, fixtures.CABLE_A, kinds, [float(t) for t in fixtures.TEMPS], prefix="E",
    )
    return built


def test_sections_split_between_anchor_structures(sections):
    assert [s.tramo_label for s in sections] == ["E5-E6", "E6-E8"]
    assert sections[0].intermediate_labels == []
    assert sections[1].intermediate_labels == ["E7"]


def test_table_10_4_values(sections):
    section = sections[0]
    assert section.ruling_span == pytest.approx(86.4359)
    assert len(section.subspans) == 1
    span = section.subspans[0]
    assert span.desnivel == pytest.approx(-17.46)
    assert round(span.vano, 1) == 88.1
    assert [round(span.sag[float(t)], 2) for t in fixtures.TEMPS] == fixtures.SAG_56
    assert [round(span.wave[float(t)], 2) for t in fixtures.TEMPS] == fixtures.WAVE_56
    expected = [74.44, 73.42, 73.42, 72.4, 72.4, 71.38, 71.38, 70.36, 70.36, 69.34, 69.34]
    assert [round(section.tension_kg[float(t)], 2) for t in fixtures.TEMPS] == expected


def test_table_10_5_values_with_intermediate_suspension(sections):
    section = sections[1]
    assert section.ruling_span == pytest.approx(83.7103)
    assert len(section.subspans) == 2
    first, second = section.subspans
    assert (first.from_label, first.to_label) == ("E6", "E7")
    assert (second.from_label, second.to_label) == ("E7", "E8")
    assert round(first.vano, 1) == 86.3
    assert round(second.vano, 1) == 80.9
    assert first.desnivel == pytest.approx(-1.44)
    assert second.desnivel == pytest.approx(-1.09)
    assert [round(first.sag[float(t)], 2) for t in fixtures.TEMPS] == fixtures.SAG_67
    assert [round(second.wave[float(t)], 2) for t in fixtures.TEMPS] == fixtures.WAVE_78
    expected = [120.33, 118.29, 115.23, 113.19, 111.15, 109.11, 107.07, 106.05, 104.01, 101.97, 100.95]
    assert [round(section.tension_kg[float(t)], 2) for t in fixtures.TEMPS] == expected


def _grid(table) -> list[list[str]]:
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


@pytest.fixture(scope="module")
def document(sections) -> Document:
    options = docx_writer.DocOptions(condicion="Initial RS", chapter="10", start_number=4)
    blob = docx_writer.build_document(sections, [float(t) for t in fixtures.TEMPS], options)
    return Document(io.BytesIO(blob))


def test_document_titles(document):
    titles = [p.text for p in document.paragraphs if p.text.startswith("Tabla ")]
    assert titles == [
        "Tabla 10-4: Tramo entre las estructuras N°5 y N°6 en condición Initial RS",
        "Tabla 10-5: Tramo entre las estructuras N°6 y N°8 en condición Initial RS",
    ]


def test_simple_table_layout(document):
    grid = _grid(document.tables[0])
    assert len(grid) == 5  # 2 de encabezado + flecha + tiempo + tension
    assert grid[1][:6] == ["Tramo", "Equivalente", "Estructuras", "Vano [m]", "[m]", "Método de tensado"]
    assert grid[1][6] == "-10°C" and grid[1][-1] == "40°C"
    assert grid[2][:6] == ["E5-E6", "86.4359", "E5-E6", "88.1", "-17.46", "Flecha en grampa [m]"]
    # Texto identico al de la tabla 10-4 de referencia (4.3, no 4.30).
    assert grid[2][6:] == [
        "4.08", "4.11", "4.14", "4.17", "4.21", "4.24", "4.27", "4.3", "4.33", "4.36", "4.39",
    ]
    assert grid[3][5] == "Tiempo [s]"
    assert grid[3][6:] == [f"{v:g}" for v in fixtures.WAVE_56]
    assert grid[4][5] == "Tensión [kg]"
    assert grid[4][6:] == [
        "74.44", "73.42", "73.42", "72.4", "72.4", "71.38", "71.38", "70.36", "70.36", "69.34", "69.34",
    ]


def test_table_with_suspension_repeats_control_block_and_single_tension_row(document):
    grid = _grid(document.tables[1])
    assert len(grid) == 7  # 2 de encabezado + 2 vanos x 2 filas + 1 de tension
    assert grid[2][:6] == ["E6-E8", "83.7103", "E6-E7", "86.3", "-1.44", "Flecha en grampa [m]"]
    assert grid[3][5] == "Tiempo [s]"
    assert grid[4][:6] == ["E6-E8", "83.7103", "E7-E8", "80.9", "-1.09", "Flecha en grampa [m]"]
    assert grid[5][5] == "Tiempo [s]"
    assert grid[6][5] == "Tensión [kg]"
    # Texto identico al de la tabla 10-5 de referencia.
    assert grid[6][6:] == [
        "120.33", "118.29", "115.23", "113.19", "111.15", "109.11",
        "107.07", "106.05", "104.01", "101.97", "100.95",
    ]
    assert grid[2][6:] == [f"{v:g}" for v in fixtures.SAG_67]
    assert grid[5][6:] == [f"{v:g}" for v in fixtures.WAVE_78]
    # La tension aparece una sola vez para todo el tramo entre anclajes.
    assert sum(1 for row in grid if row[5] == "Tensión [kg]") == 1


def _tc_at(table, row: int, column: int):
    """Devuelve (elemento tc, gridSpan) que cubre esa columna de la grilla.

    Hay que recorrer los w:tc acumulando gridSpan porque una celda combinada
    horizontalmente ocupa varias columnas con un unico elemento.
    """
    cursor = 0
    for tc in table.rows[row]._tr.findall(qn("w:tc")):
        properties = tc.find(qn("w:tcPr"))
        span_element = properties.find(qn("w:gridSpan")) if properties is not None else None
        span = int(span_element.get(qn("w:val"))) if span_element is not None else 1
        if cursor <= column < cursor + span:
            return tc, span
        cursor += span
    raise AssertionError(f"No hay celda en fila {row}, columna {column}")


def _vmerge(table, row: int, column: int) -> str | None:
    """'restart', 'continue' o None segun como este combinada verticalmente."""
    tc, _ = _tc_at(table, row, column)
    properties = tc.find(qn("w:tcPr"))
    element = properties.find(qn("w:vMerge")) if properties is not None else None
    if element is None:
        return None
    return element.get(qn("w:val")) or "continue"


def _gridspan(table, row: int, column: int) -> int:
    return _tc_at(table, row, column)[1]


def test_tramo_and_luz_equivalente_span_the_whole_section(document):
    table = document.tables[1]
    for column in (0, 1):
        assert _vmerge(table, 2, column) == "restart"
        assert [_vmerge(table, r, column) for r in range(3, 7)] == ["continue"] * 4


def test_control_block_is_merged_per_subspan(document):
    table = document.tables[1]
    for column in (2, 3, 4):
        # Un bloque de dos filas por vano; la fila de tension va aparte.
        assert [_vmerge(table, r, column) for r in range(2, 6)] == [
            "restart", "continue", "restart", "continue",
        ]
    # En la fila de tension las tres columnas de control forman una sola celda.
    assert _gridspan(table, 6, 2) == 3


def test_single_subspan_table_merges_control_block_down_to_the_tension_row(document):
    table = document.tables[0]
    for column in (2, 3, 4):
        assert [_vmerge(table, r, column) for r in range(2, 5)] == ["restart", "continue", "continue"]


def test_top_header_row_groups_control_and_temperatures(document):
    table = document.tables[0]
    assert _gridspan(table, 0, 2) == 2                      # "Control" sobre Estructuras y Vano
    assert _gridspan(table, 0, 6) == len(fixtures.TEMPS)    # franja unica sobre las temperaturas


def test_fixed_decimals_mode_keeps_trailing_zeros(sections):
    options = docx_writer.DocOptions(trim_trailing_zeros=False)
    blob = docx_writer.build_document(sections, [float(t) for t in fixtures.TEMPS], options)
    grid = _grid(Document(io.BytesIO(blob)).tables[0])
    assert grid[2][13] == "4.30"      # con recorte seria "4.3"
    assert grid[4][9] == "72.40"


def test_comma_decimal_separator(sections):
    options = docx_writer.DocOptions(decimal_separator=",")
    blob = docx_writer.build_document(sections, [float(t) for t in fixtures.TEMPS], options)
    grid = _grid(Document(io.BytesIO(blob)).tables[0])
    assert grid[2][:5] == ["E5-E6", "86,4359", "E5-E6", "88,1", "-17,46"]


def test_marking_an_intermediate_structure_as_anchor_splits_the_section(dataset):
    kinds = {k: s.auto_kind for k, s in dataset.structures.items()}
    kinds["7"] = analysis.ANCLAJE
    built, _ = analysis.build_sections(dataset, fixtures.CABLE_A, kinds, [0.0])
    assert [s.tramo_label for s in built] == ["E5-E6", "E6-E7", "E7-E8"]


def test_marking_an_endpoint_as_suspension_merges_two_sections(dataset):
    kinds = {k: s.auto_kind for k, s in dataset.structures.items()}
    kinds["6"] = analysis.SUSPENSION
    built, _ = analysis.build_sections(dataset, fixtures.CABLE_A, kinds, [0.0])
    assert [s.tramo_label for s in built] == ["E5-E8"]
    assert built[0].intermediate_labels == ["E6", "E7"]


# --------------------------------------------------------------------------
# Titulo como referencia de Word (estilo Titulo + campo SEQ)
# --------------------------------------------------------------------------
def _caption_paragraphs(document):
    return [p for p in document.paragraphs if p.style.name == "Caption"]


def test_titles_use_the_caption_style(document):
    captions = _caption_paragraphs(document)
    assert len(captions) == 2
    assert captions[0].text.startswith("Tabla 10-4: Tramo entre las estructuras")


def test_titles_carry_a_seq_field_so_word_can_cross_reference_them(document):
    instructions = [
        node.text
        for caption in _caption_paragraphs(document)
        for node in caption._p.iter(qn("w:instrText"))
    ]
    assert len(instructions) == 2
    # El primero reinicia la numeracion en el numero elegido por el usuario.
    assert instructions[0].strip() == "SEQ Tabla \\r 4 \\* ARABIC"
    assert instructions[1].strip() == "SEQ Tabla \\* ARABIC"


def test_the_number_is_cached_so_it_shows_before_pressing_f9(document):
    captions = _caption_paragraphs(document)
    assert captions[0].text == (
        "Tabla 10-4: Tramo entre las estructuras N°5 y N°6 en condición Initial RS"
    )
    assert captions[1].text.startswith("Tabla 10-5:")


def test_caption_stays_with_its_table(document):
    for caption in _caption_paragraphs(document):
        assert caption.paragraph_format.keep_with_next is True


# --------------------------------------------------------------------------
# Pagina y anchos de columna
# --------------------------------------------------------------------------
def test_default_page_is_tabloid_landscape(document):
    from docx.shared import Emu

    section = document.sections[0]
    assert round(Emu(section.page_width).cm, 1) == 43.2   # 17 pulgadas
    assert round(Emu(section.page_height).cm, 1) == 27.9  # 11 pulgadas
    assert section.page_width > section.page_height


def test_columns_are_wide_enough_for_their_widest_value(sections):
    options = docx_writer.DocOptions()
    temps = [float(t) for t in fixtures.TEMPS]
    widths = docx_writer.measure_columns(sections, temps, options)
    assert len(widths) == 6 + len(temps)
    # "Flecha en grampa [m]" es el texto mas ancho de la columna del metodo.
    assert widths[5] >= docx_writer._width_for(len("Flecha en grampa [m]"), options.font_size) - 1e-9
    # La columna de la luz equivalente debe caber "86.4359".
    assert widths[1] >= docx_writer._width_for(len("Equivalente"), options.font_size) - 1e-9


def test_widths_fit_the_printable_area(sections):
    options = docx_writer.DocOptions()
    temps = [float(t) for t in fixtures.TEMPS]
    widths = docx_writer.measure_columns(sections, temps, options)
    available = 43.18 - 2 * options.margin_cm
    assert sum(widths) <= available, "las columnas no entran en el tabloide horizontal"
    assert docx_writer.fit_widths(widths, available) == widths  # no hace falta encoger


def test_narrow_pages_shrink_the_columns_proportionally():
    widths = docx_writer.fit_widths([10.0, 10.0, 20.0], 20.0)
    assert sum(widths) == pytest.approx(20.0)
    assert widths == pytest.approx([5.0, 5.0, 10.0])


def test_table_layout_is_fixed_so_word_respects_the_widths(document):
    layout = document.tables[0]._tbl.tblPr.find(qn("w:tblLayout"))
    assert layout is not None and layout.get(qn("w:type")) == "fixed"
