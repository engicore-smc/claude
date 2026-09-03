"""Pruebas del bot de Telegram: identificacion de reportes y flujo completo."""
from __future__ import annotations

import importlib
import io
import os

import pytest
from docx import Document

from tests import fixtures

PASSWORD = "clave-bot"


@pytest.fixture(scope="module")
def bot_modules():
    os.environ["TELEGRAM_TOKEN"] = "123:fake"
    os.environ["BOT_PASSWORD"] = PASSWORD
    os.environ.pop("TELEGRAM_ALLOWED_USERS", None)
    from bot import auth, config, flow

    importlib.reload(config)
    importlib.reload(auth)
    importlib.reload(flow)
    return config, auth, flow


@pytest.fixture
def flow_mod(bot_modules):
    return bot_modules[2]


@pytest.fixture
def session(flow_mod):
    return flow_mod.store.reset(chat_id=1)


ARCHIVOS = {
    "sag": ("BD_tensado.xlsx", fixtures.sag_xlsx),
    "cable": ("flecha_tension.xlsx", fixtures.cable_xlsx),
    "structures": ("estructuras.xlsx", fixtures.structures_xlsx),
}


# --------------------------------------------------------------------------
# Identificacion
# --------------------------------------------------------------------------
def test_each_report_is_identified_by_its_columns(flow_mod):
    for esperado, (nombre, generar) in ARCHIVOS.items():
        found = flow_mod.identify(generar(), nombre)
        assert found is not None, f"no reconocio {nombre}"
        assert found[0] == esperado


def test_identification_ignores_the_file_name(flow_mod):
    # El nombre no debe influir: manda el contenido.
    found = flow_mod.identify(fixtures.structures_xlsx(), "reporte tensado final v3.xlsx")
    assert found is not None and found[0] == "structures"


def test_an_unrelated_file_is_rejected(flow_mod):
    import pandas as pd

    buffer = io.BytesIO()
    pd.DataFrame({"cosa": [1, 2], "otra": ["a", "b"]}).to_excel(buffer, index=False)
    assert flow_mod.identify(buffer.getvalue(), "cualquiera.xlsx") is None


# --------------------------------------------------------------------------
# Flujo
# --------------------------------------------------------------------------
def _cargar(flow_mod, session, keys):
    replies = []
    for key in keys:
        nombre, generar = ARCHIVOS[key]
        replies.append(flow_mod.add_report(session, generar(), nombre))
    return replies


def test_reports_can_arrive_in_any_order(flow_mod, session):
    replies = _cargar(flow_mod, session, ["structures", "cable", "sag"])
    assert "Reporte Staking table" in replies[0].text
    assert "Reporte flecha y tensión" in replies[1].text
    assert not session.faltan
    # Al completar el tercero ya ofrece los conductores.
    assert replies[2].buttons


def test_it_says_what_is_still_missing(flow_mod, session):
    reply = _cargar(flow_mod, session, ["sag"])[0]
    assert "Faltan:" in reply.text
    assert "Reporte flecha y tensión" in reply.text
    assert "Reporte Staking table" in reply.text
    assert not reply.buttons


def test_an_unknown_file_does_not_break_the_session(flow_mod, session):
    _cargar(flow_mod, session, ["sag"])
    reply = flow_mod.add_report(session, b"esto no es un xlsx", "notas.txt")
    assert "No reconocí" in reply.text
    assert session.faltan == ["cable", "structures"]  # la sesion sigue intacta


def test_resending_a_report_replaces_it(flow_mod, session):
    _cargar(flow_mod, session, ["sag", "cable", "structures"])
    reply = flow_mod.add_report(session, fixtures.sag_xlsx(), "otro_tensado.xlsx")
    assert "Reemplazo el" in reply.text
    assert session.filenames["sag"] == "otro_tensado.xlsx"


def test_cable_buttons_list_every_conductor(flow_mod, session):
    _cargar(flow_mod, session, ["sag", "cable", "structures"])
    reply = flow_mod.cable_prompt(session)
    etiquetas = [label for label, _ in reply.buttons]
    assert any(f"{fixtures.CABLE_A:g} daN/m" in e for e in etiquetas)
    assert any(f"{fixtures.CABLE_B:g} daN/m" in e for e in etiquetas)
    assert [data for _, data in reply.buttons] == ["cable:0", "cable:1"]
    assert "Initial RS" in reply.text


def test_generating_returns_the_word_document(flow_mod, session):
    _cargar(flow_mod, session, ["sag", "cable", "structures"])
    indice = next(i for i, c in enumerate(session.cables) if c["value"] == pytest.approx(fixtures.CABLE_A))
    reply = flow_mod.generate(session, indice)

    assert reply.document is not None
    nombre, blob = reply.document
    assert nombre.startswith("tablas-tensado-Initial-RS-") and nombre.endswith(".docx")

    doc = Document(io.BytesIO(blob))
    assert len(doc.tables) == 2
    titulos = [p.text for p in doc.paragraphs if p.style.name == "Caption"]
    assert titulos[0] == "Tabla 1: Tramo entre las estructuras N°5 y N°6 en condición Initial RS"
    assert titulos[1].startswith("Tabla 2:")
    assert "2 tablas" in reply.text
    assert f"{len(fixtures.TEMPS)} temperaturas" in reply.text


def test_generated_document_uses_the_defaults(flow_mod, session):
    from docx.shared import Emu

    _cargar(flow_mod, session, ["sag", "cable", "structures"])
    _, blob = flow_mod.generate(session, 0).document
    doc = Document(io.BytesIO(blob))
    page = doc.sections[0]
    assert round(Emu(page.page_width).cm, 1) == 43.2   # tabloide horizontal
    assert [c.text for c in doc.tables[0].rows[-1].cells][5] == "Tensión [kg]"


def test_generating_before_the_reports_are_in_explains_what_is_missing(flow_mod, session):
    _cargar(flow_mod, session, ["sag"])
    assert "venció" in flow_mod.generate(session, 0).text


def test_status_lists_the_three_reports(flow_mod, session):
    _cargar(flow_mod, session, ["cable"])
    texto = flow_mod.status(session).text
    assert "✅ Reporte flecha y tensión" in texto
    assert "⬜ Reporte tensado" in texto
    assert "⬜ Reporte Staking table" in texto


# --------------------------------------------------------------------------
# Acceso
# --------------------------------------------------------------------------
def test_nobody_is_authorized_until_they_give_the_password(bot_modules):
    _, auth, _ = bot_modules
    assert auth.is_authorized(999) is False
    assert auth.unlock(999, "incorrecta") is False
    assert auth.is_authorized(999) is False
    assert auth.unlock(999, PASSWORD) is True
    assert auth.is_authorized(999) is True
    auth.forget(999)
    assert auth.is_authorized(999) is False


def test_an_id_allowlist_authorizes_without_a_password(monkeypatch):
    monkeypatch.setenv("TELEGRAM_TOKEN", "123:fake")
    monkeypatch.setenv("TELEGRAM_ALLOWED_USERS", "111, 222")
    monkeypatch.delenv("BOT_PASSWORD", raising=False)
    from bot import auth, config

    importlib.reload(config)
    importlib.reload(auth)
    try:
        assert config.settings.allowed_users == {111, 222}
        assert auth.is_authorized(111) and auth.is_authorized(222)
        assert not auth.is_authorized(333)
    finally:
        importlib.reload(config)
        importlib.reload(auth)


def test_the_bot_refuses_to_start_without_any_access_control(monkeypatch):
    monkeypatch.setenv("TELEGRAM_TOKEN", "123:fake")
    monkeypatch.delenv("TELEGRAM_ALLOWED_USERS", raising=False)
    monkeypatch.delenv("BOT_PASSWORD", raising=False)
    from bot import auth, config, main

    importlib.reload(config)
    importlib.reload(auth)
    importlib.reload(main)
    try:
        with pytest.raises(SystemExit, match="TELEGRAM_ALLOWED_USERS"):
            main.build_application()
    finally:
        importlib.reload(config)
        importlib.reload(auth)
        importlib.reload(main)


def test_it_also_refuses_without_a_token(monkeypatch):
    monkeypatch.delenv("TELEGRAM_TOKEN", raising=False)
    monkeypatch.setenv("BOT_PASSWORD", "x")
    from bot import config, main

    importlib.reload(config)
    importlib.reload(main)
    try:
        with pytest.raises(SystemExit, match="TELEGRAM_TOKEN"):
            main.build_application()
    finally:
        importlib.reload(config)
        importlib.reload(main)
