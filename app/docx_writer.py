"""Generacion del anexo Word con las tablas de tensado."""
from __future__ import annotations

import io
import re
from dataclasses import dataclass

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Mm, Pt

from .analysis import DEFAULT_DECIMALS, Section

DEFAULT_TITLE_TEMPLATE = "Tabla {numero}: Tramo entre las estructuras N°{inicio} y N°{fin} en condición {condicion}"

# Etiqueta del campo SEQ: es la que Word ofrece en "Referencia cruzada".
CAPTION_LABEL = "Tabla"

ROW_LABELS = {
    "sag": "Flecha en grampa [m]",
    "wave": "Tiempo [s]",
    "tension": "Tensión [kg]",
}

# Tamanos de pagina, en vertical (ancho, alto).
PAGE_SIZES: dict[str, tuple] = {
    "tabloide": (Mm(279.4), Mm(431.8)),   # 11 x 17 pulgadas
    "a3": (Mm(297), Mm(420)),
    "a4": (Mm(210), Mm(297)),
    "carta": (Mm(215.9), Mm(279.4)),
    "oficio": (Mm(215.9), Mm(355.6)),
}

# Ancho medio de caracter (en em) y margenes internos de celda, para calcular
# columnas lo bastante anchas como para que ningun valor se parta en dos lineas.
CHAR_EM = 0.55
CELL_PADDING_CM = 0.5
PT_TO_CM = 2.54 / 72


@dataclass
class DocOptions:
    condicion: str = "Initial RS"
    title_template: str = DEFAULT_TITLE_TEMPLATE
    chapter: str = "10"
    start_number: int = 1
    font_name: str = "Calibri"
    font_size: float = 8.0
    title_size: float = 10.0
    page_size: str = "tabloide"
    landscape: bool = True
    margin_cm: float = 1.5
    decimal_separator: str = "."
    decimals: dict[str, int] | None = None
    trim_trailing_zeros: bool = True
    document_title: str = "Anexo - Tablas de tensado"
    include_document_title: bool = True

    def decimal(self, key: str) -> int:
        return {**DEFAULT_DECIMALS, **(self.decimals or {})}.get(key, 2)


def format_number(value: float | None, decimals: int, separator: str = ".", trim: bool = True) -> str:
    """Redondea a los decimales pedidos; 'trim' quita los ceros sobrantes.

    Las tablas de referencia muestran 4.3 y 72.4 en lugar de 4.30 y 72.40, que
    es como Excel imprime los numeros por defecto.
    """
    if value is None:
        return ""
    text = f"{value:.{decimals}f}"
    if trim and "." in text:
        text = text.rstrip("0").rstrip(".")
    if set(text) <= {"-", "0", "."}:
        text = text.lstrip("-") or "0"
    return text.replace(".", separator) if separator != "." else text


def format_temp(temp: float) -> str:
    return f"{int(temp)}°C" if float(temp).is_integer() else f"{temp:g}°C"


# --------------------------------------------------------------------------
# Utilidades de bajo nivel
# --------------------------------------------------------------------------
def _new(tag: str, **attrs):
    from docx.oxml import OxmlElement

    element = OxmlElement(tag)
    for key, value in attrs.items():
        element.set(qn(f"w:{key}"), value)
    return element


def _style_run(run, options: DocOptions, *, bold: bool = False, size: float | None = None) -> None:
    run.bold = bold
    run.font.size = Pt(size if size is not None else options.font_size)
    run.font.name = options.font_name
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = _new("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:eastAsia"), options.font_name)


def _write(cell, text: str, options: DocOptions, *, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """Escribe en una celda dejando un unico parrafo (util tras un merge)."""
    for extra in list(cell.paragraphs)[1:]:
        extra._element.getparent().remove(extra._element)
    paragraph = cell.paragraphs[0]
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    _style_run(paragraph.add_run(text), options, bold=bold)
    tc_pr = cell._tc.get_or_add_tcPr()
    if tc_pr.find(qn("w:vAlign")) is None:
        tc_pr.append(_new("w:vAlign", val="center"))


def _merge(table, r1: int, c1: int, r2: int, c2: int):
    return table.cell(r1, c1).merge(table.cell(r2, c2))


def _set_widths(table, widths_cm: list[float]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        tbl_pr.append(_new("w:tblLayout", type="fixed"))
    else:
        layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        for index, width in enumerate(widths_cm):
            if index < len(row.cells):
                row.cells[index].width = Cm(width)
    for index, width in enumerate(widths_cm):
        if index < len(table.columns):
            table.columns[index].width = Cm(width)


def _repeat_header(row) -> None:
    row._tr.get_or_add_trPr().append(_new("w:tblHeader", val="true"))


# --------------------------------------------------------------------------
# Ancho de columna a partir del contenido
# --------------------------------------------------------------------------
def _width_for(chars: int, font_pt: float) -> float:
    return chars * CHAR_EM * font_pt * PT_TO_CM + CELL_PADDING_CM


def measure_columns(sections: list[Section], temperatures: list[float], options: DocOptions) -> list[float]:
    """Ancho de cada columna, en cm, para que ningun texto se parta en dos lineas.

    Se mide sobre todas las tablas del documento para que salgan todas iguales.
    """
    dec = options.decimal
    sep, trim = options.decimal_separator, options.trim_trailing_zeros
    fixed = [
        {"Tramo"},
        {"Luz", "Equivalente"},
        {"Estructuras", "Control"},
        {"Vano [m]"},
        {"Desnivel", "[m]"},
        {"Método de tensado", *ROW_LABELS.values()},
    ]
    per_temp: list[set[str]] = [{format_temp(t)} for t in temperatures]

    for section in sections:
        fixed[0].add(section.tramo_label)
        fixed[1].add(format_number(section.ruling_span, dec("ruling_span"), sep, trim))
        for sub in section.subspans:
            fixed[2].add(f"{sub.from_label}-{sub.to_label}")
            fixed[3].add(format_number(sub.vano, dec("vano"), sep, trim))
            fixed[4].add(format_number(sub.desnivel, dec("desnivel"), sep, trim))
            for index, temp in enumerate(temperatures):
                per_temp[index].add(format_number(sub.sag.get(temp), dec("sag"), sep, trim))
                per_temp[index].add(format_number(sub.wave.get(temp), dec("wave"), sep, trim))
        for index, temp in enumerate(temperatures):
            per_temp[index].add(format_number(section.tension_kg.get(temp), dec("tension"), sep, trim))

    return [
        _width_for(max((len(t) for t in texts), default=1), options.font_size)
        for texts in fixed + per_temp
    ]


def fit_widths(widths: list[float], available_cm: float) -> list[float]:
    """Reduce proporcionalmente si el total no entra en el ancho util."""
    total = sum(widths)
    if total <= available_cm or total <= 0:
        return widths
    factor = available_cm / total
    return [w * factor for w in widths]


# --------------------------------------------------------------------------
# Titulo como referencia (campo SEQ dentro de un parrafo con estilo Titulo)
# --------------------------------------------------------------------------
def caption_parts(section: Section, index: int, options: DocOptions) -> tuple[str, str, str, str]:
    """Divide el titulo en (antes, prefijo del capitulo, numero, despues).

    El 'numero' es el que se emite como campo SEQ, para que Word lo numere solo
    y lo ofrezca en Insertar > Referencia cruzada.
    """
    sentinel = "\x00NUM\x00"
    start = re.search(r"(\d+)", section.from_key)
    end = re.search(r"(\d+)", section.to_key)
    rendered = options.title_template.format(
        numero=sentinel,
        inicio=start.group(1) if start else section.from_key,
        fin=end.group(1) if end else section.to_key,
        tramo=section.tramo_label,
        condicion=options.condicion,
    )
    before, _, after = rendered.partition(sentinel)
    prefix = f"{options.chapter}-" if options.chapter else ""
    return before, prefix, str(options.start_number + index), after


def _add_caption(document: Document, section: Section, index: int, options: DocOptions) -> None:
    """Titulo con estilo 'Titulo'/'Caption' y campo SEQ, como 'Insertar titulo'."""
    paragraph = document.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True

    before, prefix, number, after = caption_parts(section, index, options)
    for text in (before, prefix):
        if text:
            _style_run(paragraph.add_run(text), options, bold=True, size=options.title_size)

    # SEQ Tabla: el primero reinicia la numeracion en el numero elegido.
    instruction = f" SEQ {CAPTION_LABEL} \\* ARABIC "
    if index == 0:
        instruction = f" SEQ {CAPTION_LABEL} \\r {options.start_number} \\* ARABIC "
    _append_field(paragraph, instruction, number, options)

    if after:
        _style_run(paragraph.add_run(after), options, bold=True, size=options.title_size)


def _append_field(paragraph, instruction: str, cached: str, options: DocOptions) -> None:
    """Inserta un campo con su resultado en cache, para que se vea sin pulsar F9."""
    begin = paragraph.add_run()
    _style_run(begin, options, bold=True, size=options.title_size)
    begin._element.append(_new("w:fldChar", fldCharType="begin"))

    instr = paragraph.add_run()
    _style_run(instr, options, bold=True, size=options.title_size)
    node = _new("w:instrText")
    node.set(qn("xml:space"), "preserve")
    node.text = instruction
    instr._element.append(node)

    separate = paragraph.add_run()
    _style_run(separate, options, bold=True, size=options.title_size)
    separate._element.append(_new("w:fldChar", fldCharType="separate"))

    _style_run(paragraph.add_run(cached), options, bold=True, size=options.title_size)

    end = paragraph.add_run()
    _style_run(end, options, bold=True, size=options.title_size)
    end._element.append(_new("w:fldChar", fldCharType="end"))


# --------------------------------------------------------------------------
# Tabla
# --------------------------------------------------------------------------
def build_section_table(
    document: Document,
    section: Section,
    temperatures: list[float],
    options: DocOptions,
    widths_cm: list[float] | None = None,
):
    temp_count = len(temperatures)
    total_columns = 6 + temp_count
    subspans = section.subspans
    body_rows = 2 * len(subspans) + 1
    dec, sep, trim = options.decimal, options.decimal_separator, options.trim_trailing_zeros

    table = document.add_table(rows=2 + body_rows, cols=total_columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # --- encabezado --------------------------------------------------------
    _write(table.cell(0, 0), "", options)
    _write(table.cell(0, 1), "Luz", options, bold=True)
    _write(_merge(table, 0, 2, 0, 3), "Control", options, bold=True)
    _write(table.cell(0, 4), "Desnivel", options, bold=True)
    _write(table.cell(0, 5), "", options)
    if temp_count:
        _write(_merge(table, 0, 6, 0, 5 + temp_count), "", options)

    _write(table.cell(1, 0), "Tramo", options, bold=True)
    _write(table.cell(1, 1), "Equivalente", options, bold=True)
    _write(table.cell(1, 2), "Estructuras", options, bold=True)
    _write(table.cell(1, 3), "Vano [m]", options, bold=True)
    _write(table.cell(1, 4), "[m]", options, bold=True)
    _write(table.cell(1, 5), "Método de tensado", options, bold=True)
    for index, temp in enumerate(temperatures):
        _write(table.cell(1, 6 + index), format_temp(temp), options, bold=True)
    for row in table.rows[:2]:
        _repeat_header(row)

    # --- cuerpo ------------------------------------------------------------
    first = 2
    last = first + body_rows - 1
    _write(_merge(table, first, 0, last, 0), section.tramo_label, options)
    _write(_merge(table, first, 1, last, 1),
           format_number(section.ruling_span, dec("ruling_span"), sep, trim), options)

    for index, subspan in enumerate(subspans):
        top = first + 2 * index
        bottom = top + 1
        # Con un solo vano las columnas de control abarcan tambien la fila de tension.
        control_bottom = last if len(subspans) == 1 else bottom
        _write(_merge(table, top, 2, control_bottom, 2), f"{subspan.from_label}-{subspan.to_label}", options)
        _write(_merge(table, top, 3, control_bottom, 3),
               format_number(subspan.vano, dec("vano"), sep, trim), options)
        _write(_merge(table, top, 4, control_bottom, 4),
               format_number(subspan.desnivel, dec("desnivel"), sep, trim), options)
        _write(table.cell(top, 5), ROW_LABELS["sag"], options, align=WD_ALIGN_PARAGRAPH.LEFT)
        _write(table.cell(bottom, 5), ROW_LABELS["wave"], options, align=WD_ALIGN_PARAGRAPH.LEFT)
        for column, temp in enumerate(temperatures):
            _write(table.cell(top, 6 + column),
                   format_number(subspan.sag.get(temp), dec("sag"), sep, trim), options)
            _write(table.cell(bottom, 6 + column),
                   format_number(subspan.wave.get(temp), dec("wave"), sep, trim), options)

    if len(subspans) > 1:
        _write(_merge(table, last, 2, last, 4), "", options)
    _write(table.cell(last, 5), ROW_LABELS["tension"], options, align=WD_ALIGN_PARAGRAPH.LEFT)
    for column, temp in enumerate(temperatures):
        _write(table.cell(last, 6 + column),
               format_number(section.tension_kg.get(temp), dec("tension"), sep, trim), options)

    if widths_cm:
        _set_widths(table, widths_cm)
    return table


# --------------------------------------------------------------------------
# Documento
# --------------------------------------------------------------------------
def _setup_page(document: Document, options: DocOptions) -> float:
    page = document.sections[0]
    width, height = PAGE_SIZES.get(options.page_size, PAGE_SIZES["tabloide"])
    if options.landscape:
        page.orientation = WD_ORIENT.LANDSCAPE
        page.page_width, page.page_height = max(width, height), min(width, height)
    else:
        page.orientation = WD_ORIENT.PORTRAIT
        page.page_width, page.page_height = min(width, height), max(width, height)
    margin = Cm(options.margin_cm)
    page.left_margin = page.right_margin = page.top_margin = page.bottom_margin = margin
    return Emu(page.page_width - page.left_margin - page.right_margin).cm


def build_document(sections: list[Section], temperatures: list[float], options: DocOptions) -> bytes:
    document = Document()
    available = _setup_page(document, options)

    normal = document.styles["Normal"]
    normal.font.name = options.font_name
    normal.font.size = Pt(options.font_size)

    widths = fit_widths(measure_columns(sections, temperatures, options), available)

    if options.include_document_title and options.document_title:
        heading = document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _style_run(heading.add_run(options.document_title), options, bold=True, size=options.title_size + 3)

    for index, section in enumerate(sections):
        _add_caption(document, section, index, options)
        build_section_table(document, section, temperatures, options, widths)
        document.add_paragraph()

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
